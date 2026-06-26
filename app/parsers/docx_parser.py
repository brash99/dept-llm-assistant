from pathlib import Path

from docx import Document as DocxDocument

from app.parser import Parser
from app.knowledge import document_from_text


class DOCXParser(Parser):
    name = "DOCXParser"
    supported_suffixes = {".docx"}

    def parse(self, path, root_path):
        path = Path(path)

        docx = DocxDocument(path)

        paragraphs = []
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        table_text = []
        for table in docx.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(cells))

        text_parts = paragraphs + table_text
        text = "\n\n".join(text_parts).strip()

        metadata = {
            "num_paragraphs": len(docx.paragraphs),
            "num_tables": len(docx.tables),
            "num_table_rows": len(table_text),
            "text_length": len(text),
            "quality": "good" if len(text.strip()) >= 100 else "poor",
            "is_empty": len(text.strip()) == 0,
        }

        props = docx.core_properties
        metadata["core_properties"] = {
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "keywords": props.keywords,
            "comments": props.comments,
            "category": props.category,
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
            "last_modified_by": props.last_modified_by,
        }

        title = props.title or path.stem

        return document_from_text(
            source_path=path,
            root_path=root_path,
            text=text,
            parser=self.name,
            title=title,
            metadata=metadata,
        )
